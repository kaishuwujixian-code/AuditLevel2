import argparse
import json
import os
import re
import sys
import zipfile
from typing import List, Set
from xml.etree import ElementTree

import tkinter as tk
from tkinter import filedialog


PLACEHOLDER_PATTERN = re.compile(r"\{[^{}]+\}")
WORD_NAMESPACE = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
import json
import re
import sys
import zipfile
from typing import Iterable, List, Set
from xml.etree import ElementTree


PLACEHOLDER_PATTERN = re.compile(r"\{[^{}]+\}")
WORD_NAMESPACE = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
from typing import Iterable, List, Set

from docx import Document


PLACEHOLDER_PATTERN = re.compile(r"\{[^{}]+\}")


def _extract_from_text(text: str) -> List[str]:
    return PLACEHOLDER_PATTERN.findall(text)


def _extract_from_xml(xml_bytes: bytes) -> List[str]:
    placeholders: List[str] = []
    root = ElementTree.fromstring(xml_bytes)
    for paragraph in root.findall(".//w:p", WORD_NAMESPACE):
        text_parts = [node.text or "" for node in paragraph.findall(".//w:t", WORD_NAMESPACE)]
        if not text_parts:
            continue
        text = "".join(text_parts)
def _extract_from_paragraphs(paragraphs: Iterable) -> List[str]:
    placeholders: List[str] = []
    for paragraph in paragraphs:
        text = "".join(run.text for run in paragraph.runs)
        placeholders.extend(_extract_from_text(text))
    return placeholders


def _extract_from_tables(tables: Iterable) -> List[str]:
    placeholders: List[str] = []
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                placeholders.extend(_extract_from_paragraphs(cell.paragraphs))
                placeholders.extend(_extract_from_tables(cell.tables))
    return placeholders


def _extract_from_section(section) -> List[str]:
    placeholders: List[str] = []
    placeholders.extend(_extract_from_paragraphs(section.header.paragraphs))
    placeholders.extend(_extract_from_tables(section.header.tables))
    placeholders.extend(_extract_from_paragraphs(section.footer.paragraphs))
    placeholders.extend(_extract_from_tables(section.footer.tables))
    return placeholders


def _block_placeholders(placeholders: Iterable[str]) -> List[str]:
    blocks: List[str] = []
    for placeholder in placeholders:
        inner = placeholder.strip("{}").strip()
        if inner.lower().find("block") != -1 or inner in {"MEASURE_BLOCK", "MEASURE_SUMMARY_ROW"}:
            blocks.append(placeholder)
    return blocks


def extract_placeholders(docx_path: str) -> dict:
    placeholders: Set[str] = set()
    with zipfile.ZipFile(docx_path) as archive:
        for name in archive.namelist():
            if not name.startswith("word/"):
                continue
            if name == "word/document.xml" or name.startswith("word/header") or name.startswith("word/footer"):
                xml_bytes = archive.read(name)
                placeholders.update(_extract_from_xml(xml_bytes))
    doc = Document(docx_path)
    placeholders: Set[str] = set()

    placeholders.update(_extract_from_paragraphs(doc.paragraphs))
    placeholders.update(_extract_from_tables(doc.tables))

    for section in doc.sections:
        placeholders.update(_extract_from_section(section))

    if not placeholders:
        raise ValueError(f"No placeholders found in {docx_path}.")

    sorted_placeholders = sorted(placeholders)
    return {"placeholders": sorted_placeholders}


def _select_docx_path(default_dir: str) -> str:
    root = tk.Tk()
    root.withdraw()
    path = filedialog.askopenfilename(
        title="Select Word template",
        initialdir=default_dir,
        filetypes=[("Word document", "*.docx")],
    )
    root.destroy()
    return path


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Extract placeholders from a Word template.")
    parser.add_argument(
        "--docx",
        default="templates/level1.docx",
        help="Path to the Word template",
    )
    parser.add_argument(
        "--ui",
        action="store_true",
        help="Open a file dialog to select the Word template",
    )
    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()
    docx_path = args.docx
    if args.ui:
        selected = _select_docx_path(os.path.dirname(docx_path) or ".")
        if not selected:
            raise SystemExit("No template selected.")
        docx_path = selected
    blocks = sorted(set(_block_placeholders(sorted_placeholders)))
    return {"placeholders": sorted_placeholders, "blocks": blocks}


def main() -> int:
    docx_path = "templates/level1.docx"
    result = extract_placeholders(docx_path)
    json.dump(result, sys.stdout, indent=2)
    sys.stdout.write("\n")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
