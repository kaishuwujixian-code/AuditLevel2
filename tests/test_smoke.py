from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from core.project_store import load_project, save_project
from core.questionnaire import apply_answers_to_project, load_questionnaire_schema
from reporting.word_renderer import render_word


def test_project_roundtrip_and_render(tmp_path: Path) -> None:
    schema = load_questionnaire_schema("schemas/level1_questionnaire.schema.json")
    source_path = Path("projects/_example/project.json")
    project = json.loads(source_path.read_text(encoding="utf-8"))
    project["unknown_payload"] = {"keep": True}

    answers = project.get("answers", {})
    if not isinstance(answers, dict):
        answers = {}
    answers.update(
        {
            "client_name": "Acme Client",
            "site_address": "123 Main St",
            "building_name": "Test Tower",
            "architectural_condition": "fair",
            "heating_block": "Central plant observed; verify controls.",
        }
    )

    apply_answers_to_project(project, answers, schema, {"{ClientName}": "Acme Client"})
    project["selected_measures"] = ["condensing_boiler_retrofit", "bas_upgrade"]
    project["measure_overrides"] = {"bas_upgrade": "Override narrative for BAS upgrade."}
    project["checklist_selections"] = {
        "Walkthrough Findings": {
            "Safety Hazards": [
                "Combustion air openings unobstructed",
                "No storage in boiler room",
            ],
            "Opportunities": [
                "BAS scheduling optimization",
                "Hydronic balancing review",
            ],
        }
    }

    project_path = tmp_path / "project.json"
    save_project(str(project_path), project)

    reloaded = load_project(str(project_path))
    assert reloaded["unknown_payload"]["keep"] is True
    assert reloaded["answers"]["client_name"] == "Acme Client"

    output_path = tmp_path / "rendered.docx"
    render_word(
        template_path="templates/level1.docx",
        project_json_path=str(project_path),
        out_path=str(output_path),
    )
    assert output_path.exists()
    document = Document(str(output_path))
    rendered_text = "\n".join(
        paragraph.text for paragraph in _iter_document_paragraphs(document)
    )
    assert "{MEASURE_BLOCK}" not in rendered_text
    assert "{FINDINGS_BLOCK}" not in rendered_text
    assert "Measure 1 – Condensing Boiler Retrofit" in rendered_text
    assert "Existing Conditions:" in rendered_text
    assert "Override narrative for BAS upgrade." in rendered_text
    assert "Safety Hazards:" in rendered_text


def _iter_document_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)
