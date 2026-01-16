import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import json
import os
import sys


# -------------------- Measure templates -------------------- #

from core.measure_catalog import load_measure_catalog
from core.paths import DEFAULT_MEASURE_CATALOG
from reporting.word_renderer import render_word

DEPRECATION_NOTICE = (
    "DEPRECATED: reporting.level1_generator is legacy. "
    "Use reporting.word_renderer.render_word or tools/render_level1.py instead."
)

MEASURE_TEMPLATES = {}
CATEGORIES = []
CATEGORY_BY_MEASURE = {}

# 样式名：按你 Word 模板实际样式名称修改
STYLE_MEASURE_TITLE = "heanding 2"  # 这是你模板里的名字就行
STYLE_SECTION_SUB = "Normal"
STYLE_BODY = "Normal"

DEFAULT_MEASURE_TEMPLATES = {}
DEFAULT_CATEGORIES = []
DEFAULT_CATEGORY_BY_MEASURE = {}
DEFAULT_STYLES = {
    "measure_title_style": STYLE_MEASURE_TITLE,
    "section_subtitle_style": STYLE_SECTION_SUB,
    "body_style": STYLE_BODY,
}

DEFAULT_PLACEHOLDERS = {
    "measure_block_paragraph": "{MEASURE_BLOCK}",
    "measure_summary_table_row": "{MEASURE_SUMMARY_ROW}",
}
FINDINGS_PLACEHOLDER = "{FINDINGS_BLOCK}"
DEFAULT_SECTION_HEADINGS = {
    "existing_conditions_heading": "Existing Conditions",
    "retrofit_conditions_heading": "Retrofit Conditions",
}
DEFAULT_PAGINATION = {
    "page_break_between_measures": True,
    "no_page_break_after_last_measure": True,
}

DEFAULT_TEMPLATE_JSON_PATH = os.path.join(
    os.path.dirname(__file__),
    "..",
    "templates",
    "template.level1.json",
)


def load_level1_template(json_path):
    with open(json_path, "r", encoding="utf-8") as handle:
        data = json.load(handle)

    for key in ("word_template_requirements",):
        if key not in data:
            raise ValueError(f"Missing required field: {key}")

    requirements = data["word_template_requirements"]
    placeholders = requirements.get("docx_placeholders")
    styles = requirements.get("styles")
    section_headings = requirements.get("section_headings")
    pagination = requirements.get("pagination", {})

    if not isinstance(placeholders, dict):
        raise ValueError("Missing docx_placeholders configuration.")
    if not isinstance(styles, dict):
        raise ValueError("Missing styles configuration.")
    if not isinstance(section_headings, dict):
        raise ValueError("Missing section_headings configuration.")

    for key in ("measure_block_paragraph", "measure_summary_table_row"):
        if key not in placeholders:
            raise ValueError(f"Missing placeholder: {key}")
    for key in ("measure_title_style", "section_subtitle_style", "body_style"):
        if key not in styles:
            raise ValueError(f"Missing style definition: {key}")
    for key in ("existing_conditions_heading", "retrofit_conditions_heading"):
        if key not in section_headings:
            raise ValueError(f"Missing section heading: {key}")

    catalog = load_measure_catalog(DEFAULT_MEASURE_CATALOG)
    measures = catalog.measures
    categories = []
    for item in catalog.categories:
        if not isinstance(item, dict) or "tab_title" not in item or "code" not in item:
            raise ValueError("Invalid measure catalog category entry.")
        categories.append((item["tab_title"], item["code"]))

    category_by_measure = {key: measure["category"] for key, measure in measures.items()}
    overrides = data.get("category_by_measure_overrides", {})
    if overrides:
        if not isinstance(overrides, dict):
            raise ValueError("category_by_measure_overrides must be a mapping.")
        category_by_measure.update(overrides)

    checklists = data.get("checklists", {})
    if checklists is None:
        checklists = {}
    if not isinstance(checklists, dict):
        raise ValueError("checklists must be a mapping if provided.")

    return {
        "measures": measures,
        "categories": categories,
        "category_by_measure": category_by_measure,
        "styles": styles,
        "placeholders": {
            "measure_block_paragraph": placeholders.get(
                "measure_block_paragraph",
                DEFAULT_PLACEHOLDERS["measure_block_paragraph"],
            ),
            "measure_summary_table_row": placeholders.get(
                "measure_summary_table_row",
                DEFAULT_PLACEHOLDERS["measure_summary_table_row"],
            ),
        },
        "section_headings": section_headings,
        "pagination": pagination,
        "checklists": checklists,
        "legacy_key_map": catalog.legacy_key_map,
    }


def _load_fallback_template_config():
    return {
        "measures": {key: dict(value) for key, value in DEFAULT_MEASURE_TEMPLATES.items()},
        "categories": list(DEFAULT_CATEGORIES),
        "category_by_measure": dict(DEFAULT_CATEGORY_BY_MEASURE),
        "styles": dict(DEFAULT_STYLES),
        "placeholders": dict(DEFAULT_PLACEHOLDERS),
        "section_headings": dict(DEFAULT_SECTION_HEADINGS),
        "pagination": dict(DEFAULT_PAGINATION),
        "checklists": {},
        "legacy_key_map": {},
    }


try:
    _TEMPLATE_CONFIG = load_level1_template(DEFAULT_TEMPLATE_JSON_PATH)
except (OSError, ValueError, json.JSONDecodeError):
    _TEMPLATE_CONFIG = _load_fallback_template_config()

MEASURE_TEMPLATES = _TEMPLATE_CONFIG["measures"]
CATEGORIES = _TEMPLATE_CONFIG["categories"]
CATEGORY_BY_MEASURE = _TEMPLATE_CONFIG["category_by_measure"]
PLACEHOLDERS = _TEMPLATE_CONFIG["placeholders"]
SECTION_HEADINGS = _TEMPLATE_CONFIG["section_headings"]
PAGINATION = _TEMPLATE_CONFIG["pagination"]
STYLE_MEASURE_TITLE = _TEMPLATE_CONFIG["styles"].get("measure_title_style", STYLE_MEASURE_TITLE)
STYLE_SECTION_SUB = _TEMPLATE_CONFIG["styles"].get("section_subtitle_style", STYLE_SECTION_SUB)
STYLE_BODY = _TEMPLATE_CONFIG["styles"].get("body_style", STYLE_BODY)
CHECKLIST_SELECTIONS = {}
MEASURE_OVERRIDES = {}
PROJECT_ANSWERS = {}

# -------------------- Word helpers -------------------- #
def add_paragraph_after(paragraph, text="", style=None, bold=False):
    """
    在给定 paragraph 后面插入一个新段落并返回它。
    方便我们按顺序往下写，不用做倒序插入。
    """
    new_p_elm = OxmlElement("w:p")
    paragraph._element.addnext(new_p_elm)
    new_p = Paragraph(new_p_elm, paragraph._parent)

    if style:
        new_p.style = style

    if text:
        run = new_p.add_run(text)
        if bold:
            run.bold = True
    else:
        # 没有文字但需要加粗其实没意义，这里就不处理 bold 了
        pass

    return new_p


#-----------填 summary 表格的函数-------------------#
def fill_measure_summary_table(doc, selected_keys):
    """
    在 Word 模板里找到包含 {MEASURE_SUMMARY_ROW} 的那一行，
    用选中的 measures 填充 summary 表格。

    假设：
    - 这一行所在表格有两列：
        第 1 列：Description of Measure
        第 2 列：Estimated Utility / Cost Savings
    - 这一行的第 1 个单元格内容是 {MEASURE_SUMMARY_ROW}
    """
    placeholder = PLACEHOLDERS.get(
        "measure_summary_table_row",
        DEFAULT_PLACEHOLDERS["measure_summary_table_row"],
    )
    target_table = None
    target_row = None

    # 找到带占位符的那一行
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    target_table = table
                    target_row = row
                    break
            if target_row is not None:
                break
        if target_row is not None:
            break

    # 如果模板里没有这个表格，就直接跳过，不报错
    if target_row is None or target_table is None:
        return

    # 如果没有选 measure，就把这一行清空
    if not selected_keys:
        for cell in target_row.cells:
            cell.text = ""
        return

    first = True
    for key in selected_keys:
        tpl = MEASURE_TEMPLATES.get(key) or {}
        name = tpl.get("name", key)
        summary = tpl.get("summary", "")

        if first:
            row = target_row
            first = False
        else:
            row = target_table.add_row()

        # 假设表格只有两列；如果有第三列可以在这里加
        row.cells[0].text = name
        row.cells[1].text = summary

    # 确保占位符被清掉
    if placeholder in row.cells[0].text:
        row.cells[0].text = row.cells[0].text.replace(placeholder, "")

# -------------------- Measure insert -------------------- #
def insert_findings_block(doc):
    """
    Insert narrative paragraphs for each checklist category into the {FINDINGS_BLOCK} placeholder.
    Each category (e.g., Operations & Maintenance, Opportunities, Safety Hazards) becomes one paragraph.
    """
    placeholder = FINDINGS_PLACEHOLDER
    anchor_idx = None
    for i, para in enumerate(doc.paragraphs):
        if placeholder in para.text:
            anchor_idx = i
            break
    if anchor_idx is None:
        return

    anchor_para = doc.paragraphs[anchor_idx]
    anchor_para.text = ""
    try:
        anchor_para.style = STYLE_BODY
    except Exception:
        pass

    current_para = anchor_para
    categories = CHECKLIST_SELECTIONS.get("Walkthrough Findings", {})
    first_para_done = False
    for category, items in categories.items():
        if not items or len(items) == 0:
            continue
        text_items = [item for item in items]
        if len(text_items) == 1:
            items_str = text_items[0]
        elif len(text_items) == 2:
            items_str = f"{text_items[0]} and {text_items[1]}"
        else:
            items_str = ", ".join(text_items[:-1]) + ", and " + text_items[-1]
        paragraph_text = f"{category}: {items_str}."
        if not first_para_done:
            current_para.text = paragraph_text
            first_para_done = True
        else:
            current_para = add_paragraph_after(current_para, paragraph_text, STYLE_BODY, bold=False)

    if not first_para_done:
        anchor_para.text = ""


def insert_measures_into_docx(template_path, output_path, selected_keys):
    """
    在 template_path 的 docx 中寻找 {MEASURE_BLOCK} 所在段落，
    按顺序把选中的 measures 写进去，然后保存到 output_path。
    - 标题用 STYLE_MEASURE_TITLE
    - Existing / Retrofit 副标题用 STYLE_SECTION_SUB 并加粗
    - 正文用 STYLE_BODY
    - 每个 measure 的最后插一个分页符（下一条新起一页），最后一条不分页
    Also fills the measures summary table and checklist findings if present.
    """
    doc = Document(template_path)

    override_text = PROJECT_ANSWERS.get("measures_block_override") or ""
    if override_text.strip():
        fill_measure_summary_table(doc, selected_keys)
        placeholder = PLACEHOLDERS.get(
            "measure_block_paragraph",
            DEFAULT_PLACEHOLDERS["measure_block_paragraph"],
        )
        anchor_idx = None
        for i, para in enumerate(doc.paragraphs):
            if placeholder in para.text:
                anchor_idx = i
                break
        if anchor_idx is not None:
            anchor_para = doc.paragraphs[anchor_idx]
            anchor_para.text = ""
            try:
                anchor_para.style = STYLE_BODY
            except Exception:
                pass
            override_parts = [p for p in override_text.split("\n\n") if p.strip() != ""]
            if len(override_parts) == 0:
                override_parts = [override_text]
            current_para = anchor_para
            first_part = True
            for part in override_parts:
                text = part.strip()
                if first_part:
                    current_para.text = text
                    first_part = False
                else:
                    current_para = add_paragraph_after(current_para, text, STYLE_BODY, bold=False)
            if placeholder in current_para.text:
                current_para.text = current_para.text.replace(placeholder, "")
        insert_findings_block(doc)
        doc.save(output_path)
        return

    fill_measure_summary_table(doc, selected_keys)

    placeholder = PLACEHOLDERS.get(
        "measure_block_paragraph",
        DEFAULT_PLACEHOLDERS["measure_block_paragraph"],
    )
    anchor_idx = None
    for i, para in enumerate(doc.paragraphs):
        if placeholder in para.text:
            anchor_idx = i
            break

    if anchor_idx is None:
        raise RuntimeError("Placeholder {MEASURE_BLOCK} not found in template.")

    anchor_para = doc.paragraphs[anchor_idx]
    anchor_para.text = ""
    if STYLE_MEASURE_TITLE:
        anchor_para.style = STYLE_MEASURE_TITLE

    current_para = anchor_para
    total = len(selected_keys)

    if total == 0:
        doc.save(output_path)
        return

    for idx, key in enumerate(selected_keys, start=1):
        tpl = MEASURE_TEMPLATES.get(key) or {}
        measure_name = tpl.get("name", str(key))

        title_text = f"3.{idx} Measure – {measure_name}"
        if idx == 1:
            current_para.add_run(title_text)
        else:
            current_para = add_paragraph_after(current_para, "", STYLE_BODY, bold=False)
            current_para = add_paragraph_after(current_para, title_text, STYLE_MEASURE_TITLE, bold=False)

        override_para_text = None
        if key in MEASURE_OVERRIDES and MEASURE_OVERRIDES[key].strip():
            override_para_text = MEASURE_OVERRIDES[key].strip()
        if override_para_text:
            override_parts = [p for p in override_para_text.split("\n\n") if p.strip() != ""]
            if len(override_parts) == 0:
                override_parts = [override_para_text]
            for part in override_parts:
                text = part.strip()
                current_para = add_paragraph_after(current_para, text, STYLE_BODY, bold=False)
        else:
            exist_heading = SECTION_HEADINGS.get(
                "existing_conditions_heading",
                DEFAULT_SECTION_HEADINGS["existing_conditions_heading"],
            )
            retro_heading = SECTION_HEADINGS.get(
                "retrofit_conditions_heading",
                DEFAULT_SECTION_HEADINGS["retrofit_conditions_heading"],
            )
            exist_sub = add_paragraph_after(current_para, exist_heading, STYLE_SECTION_SUB, bold=True)
            exist_body_text = tpl.get("existing", "")
            exist_body = add_paragraph_after(exist_sub, exist_body_text, STYLE_BODY, bold=False)
            blank_line = add_paragraph_after(exist_body, "", STYLE_BODY, bold=False)
            retro_sub = add_paragraph_after(blank_line, retro_heading, STYLE_SECTION_SUB, bold=True)
            retro_body_text = tpl.get("retrofit", "")
            retro_body = add_paragraph_after(retro_sub, retro_body_text, STYLE_BODY, bold=False)
            current_para = retro_body

        end_blank = add_paragraph_after(current_para, "", STYLE_BODY, bold=False)
        current_para = end_blank

        if idx != total and PAGINATION.get("page_break_between_measures", True):
            pb_para = add_paragraph_after(current_para, "", STYLE_BODY, bold=False)
            pb_run = pb_para.add_run()
            try:
                pb_run.add_break(WD_BREAK.PAGE)
            except NameError:
                pass
            current_para = pb_para

    insert_findings_block(doc)
    doc.save(output_path)



# -------------------- GUI Application -------------------- #

class MeasureToWordApp:
    def __init__(self, root):
        self.root = root
        root.title("Measure → Word Report Generator")
        root.geometry("800x500")

        self.template_path = tk.StringVar()
        self.output_path = tk.StringVar()

        # 上：模板选择
        top = ttk.Frame(root, padding=8)
        top.pack(side=tk.TOP, fill=tk.X)

        ttk.Label(top, text="Word Template:").pack(side=tk.LEFT)
        ttk.Entry(top, textvariable=self.template_path, width=60).pack(side=tk.LEFT, padx=5)
        ttk.Button(top, text="Browse...", command=self.browse_template).pack(side=tk.LEFT)

        # 中：左边 measure 选择，右边预览
        # 中：左边 measure 选择（分 TAB），右边预览
        main = ttk.PanedWindow(root, orient=tk.HORIZONTAL)
        main.pack(fill=tk.BOTH, expand=True, padx=8, pady=8)

        # 左侧整体 panel（里面上半 Notebook，下半按钮）
        left_panel = ttk.Frame(main)
        main.add(left_panel, weight=1)

        # --- Notebook: 按类别分 TAB ---
        nb = ttk.Notebook(left_panel)
        nb.pack(fill=tk.BOTH, expand=True, pady=(0, 5))

        # 每个类别一个 frame
        self.measure_vars = {}
        tab_frames = {}
        row_index_by_cat = {}

        for tab_title, cat_code in CATEGORIES:
            frame = ttk.Labelframe(nb, text=tab_title)
            nb.add(frame, text=tab_title)
            tab_frames[cat_code] = frame
            row_index_by_cat[cat_code] = 0

        # 把每个 measure 放进对应的 tab（不在映射里的归到 "other"）
        for key, measure in MEASURE_TEMPLATES.items():
            cat_code = CATEGORY_BY_MEASURE.get(key, "other")
            frame = tab_frames.get(cat_code, tab_frames["other"])

            row = row_index_by_cat[cat_code]
            var = tk.BooleanVar(value=False)
            self.measure_vars[key] = var
            label = measure.get("name") or measure.get("title") or key
            ttk.Checkbutton(frame, text=label, variable=var).grid(
                row=row, column=0, sticky="w", padx=4, pady=2
            )
            row_index_by_cat[cat_code] += 1

        # --- 下方：Select All / Clear 按钮（对全部 measure 生效）---
        btn_frame = ttk.Frame(left_panel)
        btn_frame.pack(side=tk.BOTTOM, anchor="w", pady=(5, 0))
        ttk.Button(btn_frame, text="Select All", command=self.select_all).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="Clear", command=self.clear_all).pack(side=tk.LEFT, padx=2)

        # 右侧预览保持原样
        right = ttk.Labelframe(main, text="Preview (plain text)")
        main.add(right, weight=2)

        self.preview_text = tk.Text(right, wrap="word")
        self.preview_text.pack(fill=tk.BOTH, expand=True, padx=4, pady=4)


        

        # 下：生成按钮
        bottom = ttk.Frame(root, padding=8)
        bottom.pack(side=tk.BOTTOM, fill=tk.X)

        ttk.Button(bottom, text="Preview Text", command=self.update_preview).pack(side=tk.LEFT)
        ttk.Button(bottom, text="Generate Word Report", command=self.generate_word).pack(side=tk.RIGHT)

    # ------ UI helpers ------ #

    def browse_template(self):
        path = filedialog.askopenfilename(
            title="Select Word Template",
            filetypes=[("Word files", "*.docx"), ("All files", "*.*")]
        )
        if path:
            self.template_path.set(path)

    def select_all(self):
        for v in self.measure_vars.values():
            v.set(True)

    def clear_all(self):
        for v in self.measure_vars.values():
            v.set(False)

    def get_selected_keys(self):
        return [k for k, v in self.measure_vars.items() if v.get()]

    # ------ Preview / Generate ------ #

    def update_preview(self):
        self.preview_text.delete("1.0", tk.END)
        selected = self.get_selected_keys()
        if not selected:
            self.preview_text.insert(tk.END, "No measures selected.\n")
            return

        existing_heading = SECTION_HEADINGS.get(
            "existing_conditions_heading",
            DEFAULT_SECTION_HEADINGS["existing_conditions_heading"],
        )
        retrofit_heading = SECTION_HEADINGS.get(
            "retrofit_conditions_heading",
            DEFAULT_SECTION_HEADINGS["retrofit_conditions_heading"],
        )

        idx = 1
        for key in selected:
            tpl = MEASURE_TEMPLATES[key]
            self.preview_text.insert(tk.END, f"3.{idx} Measure – {tpl['name']}\n\n")
            self.preview_text.insert(tk.END, f"{existing_heading}\n")
            self.preview_text.insert(tk.END, tpl["existing"] + "\n\n")
            self.preview_text.insert(tk.END, f"{retrofit_heading}\n")
            self.preview_text.insert(tk.END, tpl["retrofit"] + "\n\n")
            self.preview_text.insert(tk.END, "-" * 70 + "\n\n")
            idx += 1

    def generate_word(self):
        template = self.template_path.get().strip()
        if not template:
            messagebox.showerror("Error", "Please select a Word template first.")
            return
        if not os.path.exists(template):
            messagebox.showerror("Error", "Template file not found.")
            return

        selected = self.get_selected_keys()
        if not selected:
            messagebox.showerror("Error", "Please select at least one measure.")
            return

        out_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word files", "*.docx")],
            title="Save generated report as"
        )
        if not out_path:
            return

        try:
            insert_measures_into_docx(template, out_path, selected)
            messagebox.showinfo("Success", f"Report generated:\n{out_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate report:\n{e}")


if __name__ == "__main__":
    _warn_deprecated()
    root = tk.Tk()
    try:
        style = ttk.Style()
        if "vista" in style.theme_names():
            style.theme_use("vista")
    except Exception:
        pass

    app = MeasureToWordApp(root)
    root.mainloop()



def generate_level1_report(
    project_json_path: str,
    template_json_path: str,
    docx_template_path: str,
    out_path: str,
) -> str:
    """
    Headless entrypoint (no Tkinter).
    Deprecated: use reporting.word_renderer.render_word instead.
    The template_json_path argument is ignored and kept for compatibility.
    """
    render_word(
        template_path=docx_template_path,
        project_json_path=project_json_path,
        out_path=out_path,
    )
    return out_path


def _warn_deprecated() -> None:
    print(DEPRECATION_NOTICE, file=sys.stderr)
