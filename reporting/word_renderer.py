import json
import os
import re
import zipfile
from typing import Any, Dict, Iterable, List, Optional, Set, Tuple

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph

from core.measure_catalog import load_measure_catalog
from core.project_store import normalize_measures_data
from reporting.narratives import facility_overview, measures as measure_narratives
from reporting.narratives.registry import KNOWN_BLOCK_PLACEHOLDERS, get_block_renderer


PLACEHOLDER_PATTERN = re.compile(r"\{[^{}]+\}")
DEFAULT_MAPPING_PATH = os.path.join("schemas", "level1_placeholders.map.json")
DEFAULT_EMPTY_BLOCK_TEXT = ""
MEASURE_CATEGORY_LABELS = {
    "bas": "BAS / Controls",
    "boiler": "Boiler / Plant",
    "boilers": "Boilers",
    "dhw": "DHW",
    "lighting": "Lighting",
    "ventilation": "Ventilation",
    "mua": "MUA / Ventilation",
    "controls": "Controls",
    "loop": "Hydronic Loops",
    "water": "Water & DHW",
    "pumps": "Pumps / Power / PF",
    "envelope": "Building Envelope",
    "other": "Other / Misc",
}


def _normalize_key(text: str) -> str:
    normalized = re.sub(r"[^a-z0-9]+", "_", text.strip().lower())
    return normalized.strip("_")


def _iter_paragraphs_in_tables(tables: Iterable) -> Iterable:
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
                for paragraph in _iter_paragraphs_in_tables(cell.tables):
                    yield paragraph


def _iter_all_paragraphs(doc: Document) -> Iterable:
    for paragraph in doc.paragraphs:
        yield paragraph
    for paragraph in _iter_paragraphs_in_tables(doc.tables):
        yield paragraph
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for paragraph in _iter_paragraphs_in_tables(section.header.tables):
            yield paragraph
        for paragraph in section.footer.paragraphs:
            yield paragraph
        for paragraph in _iter_paragraphs_in_tables(section.footer.tables):
            yield paragraph


def _stringify_value(value: Any) -> Optional[str]:
    if value is None:
        return None
    if isinstance(value, (list, tuple)):
        return ", ".join(str(item) for item in value)
    if isinstance(value, dict):
        return json.dumps(value)
    return str(value)


def _has_meaningful_value(value: Any) -> bool:
    string_value = _stringify_value(value)
    return string_value is not None and bool(string_value.strip())


def _is_block_placeholder(placeholder: str) -> bool:
    if not placeholder.startswith("{") or not placeholder.endswith("}"):
        return False
    inner_text = placeholder[1:-1].strip()
    return (
        inner_text.endswith(" block")
        or inner_text.endswith("_BLOCK")
        or placeholder in KNOWN_BLOCK_PLACEHOLDERS
    )


def _load_mapping(mapping_path: Optional[str]) -> Optional[Dict[str, str]]:
    if mapping_path is None:
        mapping_path = DEFAULT_MAPPING_PATH if os.path.isfile(DEFAULT_MAPPING_PATH) else None
    if mapping_path is None:
        return None
    if not os.path.isfile(mapping_path):
        raise FileNotFoundError(f"Mapping file not found: {mapping_path}")
    with open(mapping_path, "r", encoding="utf-8") as handle:
        mapping_data = json.load(handle)
    if not isinstance(mapping_data, dict):
        raise ValueError("Mapping file must contain a JSON object of placeholder-to-question mappings.")
    return {str(key): str(value) for key, value in mapping_data.items()}


def _build_placeholder_map_from_placeholders(placeholders: Dict[str, Any]) -> Dict[str, str]:
    placeholder_map: Dict[str, str] = {}
    for placeholder, value in placeholders.items():
        string_value = _stringify_value(value)
        if string_value is not None:
            placeholder_map[str(placeholder)] = string_value
    return placeholder_map




def _build_placeholder_map_from_answers(
    answers: Dict[str, Any],
    placeholders: Iterable[str],
    mapping_path: Optional[str],
) -> Dict[str, str]:
    placeholder_map: Dict[str, str] = {}
    mapping_data = _load_mapping(mapping_path)

    if mapping_data:
        for placeholder, question_id in mapping_data.items():
            value = answers.get(question_id)
            string_value = _stringify_value(value)
            if string_value is not None:
                placeholder_map[placeholder] = string_value

    for question_id, value in answers.items():
        string_value = _stringify_value(value)
        if string_value is not None:
            placeholder_map.setdefault(f"{{{question_id}}}", string_value)

    if mapping_data is None:
        normalized_answers = {
            _normalize_key(question_id): value
            for question_id, value in answers.items()
            if _stringify_value(value) is not None
        }
        for placeholder in placeholders:
            if placeholder in placeholder_map:
                continue
            inner_text = placeholder[1:-1]
            normalized_placeholder = _normalize_key(inner_text)
            if normalized_placeholder in normalized_answers:
                placeholder_map[placeholder] = _stringify_value(
                    normalized_answers[normalized_placeholder]
                )

    return placeholder_map


def _replace_placeholders_in_text(text: str, placeholder_map: Dict[str, str]) -> Tuple[str, int]:
    replaced_text = text
    replacements = 0
    for placeholder, value in placeholder_map.items():
        if placeholder in replaced_text:
            replacements += replaced_text.count(placeholder)
            replaced_text = replaced_text.replace(placeholder, value)
    return replaced_text, replacements


def _replace_placeholders_in_runs(paragraph: Paragraph, placeholder_map: Dict[str, str]) -> int:
    replacements = 0
    if not paragraph.runs or "{" not in paragraph.text:
        return replacements

    for placeholder, value in placeholder_map.items():
        while True:
            runs = paragraph.runs
            full_text = "".join(run.text for run in runs)
            start = full_text.find(placeholder)
            if start == -1:
                break
            end = start + len(placeholder)
            cursor = 0
            start_run = None
            end_run = None
            start_offset = 0
            end_offset = 0
            for index, run in enumerate(runs):
                run_len = len(run.text)
                if start_run is None and cursor + run_len > start:
                    start_run = index
                    start_offset = start - cursor
                if cursor + run_len >= end:
                    end_run = index
                    end_offset = end - cursor
                    break
                cursor += run_len
            if start_run is None or end_run is None:
                break
            if start_run == end_run:
                run = runs[start_run]
                run.text = run.text[:start_offset] + value + run.text[end_offset:]
            else:
                first_run = runs[start_run]
                last_run = runs[end_run]
                prefix = first_run.text[:start_offset]
                suffix = last_run.text[end_offset:]
                first_run.text = prefix + value + suffix
                for idx in range(start_run + 1, end_run + 1):
                    runs[idx].text = ""
            replacements += 1
    return replacements


def _set_cell_text(cell, text: str) -> None:
    if not cell.paragraphs:
        cell.add_paragraph(text)
        return
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)




def _set_cell_paragraph_style_and_alignment(cell, *, style: str | None = None, alignment: int | None = None) -> None:
    if not cell.paragraphs:
        cell.add_paragraph("")
    paragraph = cell.paragraphs[0]
    if style:
        try:
            paragraph.style = style
        except Exception:
            pass
    if alignment is not None:
        paragraph.alignment = alignment


def _set_cell_width(cell, width_inches: float) -> None:
    try:
        cell.width = Inches(width_inches)
    except Exception:
        return


def _add_paragraph_after(paragraph: Paragraph, text: str = "", style=None) -> Paragraph:
    new_p_elm = OxmlElement("w:p")
    paragraph._element.addnext(new_p_elm)
    new_p = Paragraph(new_p_elm, paragraph._parent)
    if style is not None:
        new_p.style = style
    if text:
        new_p.add_run(text)
    return new_p


def _ensure_paragraph_style(
    doc: Document, name: str, *, base: Optional[str] = None, bold: bool = False
) -> str:
    try:
        style = doc.styles[name]
        created = False
    except KeyError:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        created = True
        if base and base in doc.styles:
            style.base_style = doc.styles[base]
    if created:
        style.font.bold = bold
    return name


def _ensure_measure_styles(doc: Document) -> Dict[str, str]:
    if "Content A" in doc.styles:
        body_style = "Content A"
    else:
        body_style = _ensure_paragraph_style(doc, "Body", base="Normal", bold=False)
    subtitle_style = _ensure_paragraph_style(doc, "Section Subtitle", base=body_style, bold=True)
    title_style = "Heading 2" if "Heading 2" in doc.styles else "Heading 3"
    return {"body": body_style, "subtitle": subtitle_style, "title": title_style}


def _split_text_lines(text: str) -> List[str]:
    return [line.strip() for line in str(text).splitlines() if line.strip()]


def _format_key_inputs(measure: Dict[str, Any]) -> List[str]:
    fields = [
        ("Category", _label_for_category(measure.get("category"))),
        ("Electric savings (kWh)", measure.get("savings_electric_kwh")),
        ("Gas savings (m³)", measure.get("savings_gas_m3")),
        ("Water savings (m³)", measure.get("savings_water_m3")),
        ("Implementation cost", measure.get("implementation_cost")),
        ("Incentive", measure.get("incentive")),
        ("Simple payback (yrs)", measure.get("simple_payback_years")),
    ]
    parts: List[str] = []
    for label, value in fields:
        formatted = _format_numeric_value(value)
        if formatted is None:
            continue
        parts.append(f"{label}: {formatted}")
    return parts


def _format_numeric_value(value: Any) -> Optional[str]:
    if value is None:
        return None
    if isinstance(value, str):
        if not value.strip():
            return None
        return value.strip()
    if isinstance(value, float):
        return f"{value:,.2f}".rstrip("0").rstrip(".")
    if isinstance(value, int):
        return f"{value:,d}"
    return str(value)


def _label_for_category(value: Any) -> Optional[str]:
    if value is None:
        return None
    text = str(value).strip()
    if not text:
        return None
    return MEASURE_CATEGORY_LABELS.get(text.lower(), text)


def _format_measure_heading(index: int, title: str) -> str:
    clean_title = title.strip() if isinstance(title, str) else ""
    if not clean_title:
        clean_title = "Measure"
    return f"Measure {index} – {clean_title}"


def _insert_measure_block(
    paragraph: Paragraph,
    measures: List[Dict[str, Any]],
    styles: Dict[str, str],
    fallback_text: str,
) -> None:
    def set_paragraph(target: Paragraph, text: str, style: str) -> Paragraph:
        target.text = ""
        target.style = style
        if text:
            target.add_run(text)
        return target

    if not measures:
        set_paragraph(paragraph, fallback_text, styles["body"])
        return

    current_para = paragraph
    first = True
    total = len(measures)
    for index, measure in enumerate(measures, start=1):
        title = _format_measure_heading(index, str(measure.get("measure_title", "")).strip())
        if first:
            current_para = set_paragraph(current_para, title, styles["title"])
            first = False
        else:
            current_para = _add_paragraph_after(current_para, title, style=styles["title"])

        existing = measure.get("existing_conditions")
        if existing:
            current_para = _add_paragraph_after(
                current_para, "Existing Conditions:", style=styles["subtitle"]
            )
            for line in _split_text_lines(existing):
                current_para = _add_paragraph_after(current_para, line, style=styles["body"])
                current_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        retrofit = measure.get("retrofit_conditions")
        if retrofit:
            current_para = _add_paragraph_after(
                current_para, "Retrofit Conditions:", style=styles["subtitle"]
            )
            for line in _split_text_lines(retrofit):
                current_para = _add_paragraph_after(current_para, line, style=styles["body"])
                current_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        key_inputs = _format_key_inputs(measure)
        if key_inputs:
            current_para = _add_paragraph_after(current_para, "Key Inputs:", style=styles["subtitle"])
            current_para = _add_paragraph_after(
                current_para, "; ".join(key_inputs), style=styles["body"]
            )
            current_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Notes are summarized in {MEASURE_SUMMARY_ROW}; avoid repeating here.
        if index < total:
            current_para = _add_paragraph_after(current_para, "", style=styles["body"])
            pb_run = current_para.add_run()
            pb_run.add_break(WD_BREAK.PAGE)


def _resolve_measure_summary_text(measure: Dict[str, Any]) -> str:
    for key in ("notes", "summary", "retrofit_conditions", "existing_conditions"):
        value = measure.get(key)
        if value is None:
            continue
        text = str(value).strip()
        if text:
            return text
    return ""


def _collect_measure_summary_rows(
    project_data: Dict[str, Any],
    measures: List[Dict[str, Any]],
    catalog: Any,
) -> List[Tuple[str, str]]:
    rows: List[Tuple[str, str]] = []
    if measures:
        for measure in measures:
            title = str(measure.get("measure_title", "")).strip() or "Measure"
            summary = _resolve_measure_summary_text(measure)
            rows.append((title, summary))
        return rows

    if catalog is None:
        return rows

    selected_ids = measure_narratives.collect_selected_measure_ids(project_data, catalog=catalog)
    for measure_id in selected_ids:
        catalog_measure = catalog.measures.get(measure_id, {})
        title = (
            str(catalog_measure.get("title", "")).strip()
            or str(catalog_measure.get("name", "")).strip()
            or measure_id
        )
        summary = str(catalog_measure.get("summary", "")).strip()
        rows.append((title, summary))
    return rows


def _fill_measure_summary_table(
    doc: Document,
    project_data: Dict[str, Any],
    measures: List[Dict[str, Any]],
    catalog: Any,
) -> bool:
    placeholder = "{MEASURE_SUMMARY_ROW}"
    target_table = None
    target_row = None
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    target_table = table
                    target_row = row
                    break
            if target_row is not None:
                break
        if target_row is not None:
            break

    if target_row is None or target_table is None:
        return False

    summary_rows = _collect_measure_summary_rows(project_data, measures, catalog)
    if not summary_rows:
        for cell in target_row.cells:
            _set_cell_text(cell, "")
        return True

    body_style = "Content A" if "Content A" in doc.styles else None
    for index, (title, summary) in enumerate(summary_rows, start=1):
        row = target_row if index == 1 else target_table.add_row()
        if len(row.cells) >= 3:
            _set_cell_text(row.cells[0], str(index))
            _set_cell_paragraph_style_and_alignment(
                row.cells[0], style=body_style, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
            )
            _set_cell_width(row.cells[0], 0.65)
            _set_cell_text(row.cells[1], title)
            _set_cell_paragraph_style_and_alignment(row.cells[1], style=body_style)
            _set_cell_text(row.cells[2], summary)
            _set_cell_paragraph_style_and_alignment(row.cells[2], style=body_style)
        elif len(row.cells) == 2:
            _set_cell_text(row.cells[0], f"Measure {index} – {title}")
            _set_cell_paragraph_style_and_alignment(row.cells[0], style=body_style)
            _set_cell_text(row.cells[1], summary)
            _set_cell_paragraph_style_and_alignment(row.cells[1], style=body_style)
        elif len(row.cells) == 1:
            _set_cell_text(
                row.cells[0],
                f"Measure {index} – {title} — {summary}" if summary else f"Measure {index} – {title}",
            )
            _set_cell_paragraph_style_and_alignment(row.cells[0], style=body_style)

    for cell in target_row.cells:
        if placeholder in cell.text:
            _set_cell_text(cell, cell.text.replace(placeholder, ""))
    return True


def _split_block_paragraphs(text: str) -> List[str]:
    lines = text.splitlines()
    paragraphs: List[str] = []
    buffer: List[str] = []
    for line in lines:
        cleaned = line.strip()
        if not cleaned:
            if buffer:
                paragraphs.append(" ".join(buffer).strip())
                buffer = []
            continue
        buffer.append(cleaned)
    if buffer:
        paragraphs.append(" ".join(buffer).strip())
    if not paragraphs and text.strip():
        return [text.strip()]
    return paragraphs


def _iter_xml_parts(docx_path: str) -> Iterable[Tuple[str, bytes]]:
    with zipfile.ZipFile(docx_path) as archive:
        for info in archive.infolist():
            if not info.filename.startswith("word/") or not info.filename.endswith(".xml"):
                continue
            if info.filename == "word/document.xml" or info.filename.startswith(
                ("word/header", "word/footer")
            ):
                yield info.filename, archive.read(info.filename)


def _collect_placeholders_from_docx(docx_path: str) -> List[str]:
    placeholders: List[str] = []
    for _, xml_bytes in _iter_xml_parts(docx_path):
        text = xml_bytes.decode("utf-8", errors="ignore")
        placeholders.extend(PLACEHOLDER_PATTERN.findall(text))
    return placeholders


def _replace_placeholders_in_docx_xml(
    docx_path: str, placeholder_map: Dict[str, str]
) -> int:
    if not placeholder_map:
        return 0
    temp_path = f"{docx_path}.tmp"
    replacements = 0
    with zipfile.ZipFile(docx_path, "r") as archive, zipfile.ZipFile(
        temp_path, "w", compression=zipfile.ZIP_DEFLATED
    ) as output:
        for info in archive.infolist():
            data = archive.read(info.filename)
            if info.filename.startswith("word/") and info.filename.endswith(".xml") and (
                info.filename == "word/document.xml"
                or info.filename.startswith(("word/header", "word/footer"))
            ):
                text = data.decode("utf-8", errors="ignore")
                replaced_text, part_replacements = _replace_placeholders_in_text(
                    text, placeholder_map
                )
                replacements += part_replacements
                data = replaced_text.encode("utf-8")
            output.writestr(info, data)
    os.replace(temp_path, docx_path)
    return replacements


def render_word(
    template_path: str,
    project_json_path: str,
    out_path: str,
    mapping_path: Optional[str] = None,
    strict: bool = False,
) -> Dict[str, Any]:
    """
    Render a Word document by replacing placeholders using project JSON.

    Note: python-docx does not expose text inside shapes/textboxes, so those placeholders
    may remain unresolved until the XML fallback replacement runs.
    """
    with open(project_json_path, "r", encoding="utf-8") as handle:
        project_data = json.load(handle)
    if isinstance(project_data, dict):
        normalize_measures_data(project_data)

    doc = Document(template_path)
    placeholder_occurrences = _collect_placeholders_from_docx(template_path)
    placeholder_set = set(placeholder_occurrences)

    project_placeholders = project_data.get("placeholders")
    answers = project_data.get("answers", {})
    if "placeholders" in project_data and isinstance(project_placeholders, dict):
        placeholder_map = _build_placeholder_map_from_placeholders(project_placeholders)
        facility_overview.apply_facility_placeholders(project_data, placeholder_map)
    else:
        if not isinstance(answers, dict):
            raise ValueError("project['answers'] must be a JSON object.")
        placeholder_map = _build_placeholder_map_from_answers(
            answers, placeholder_set, mapping_path
        )
        facility_overview.apply_facility_placeholders(project_data, placeholder_map)

    mapping_data = _load_mapping(mapping_path)
    base_placeholder_map = dict(placeholder_map)
    block_placeholders = sorted(
        {
            *[ph for ph in placeholder_occurrences if _is_block_placeholder(ph)],
            *KNOWN_BLOCK_PLACEHOLDERS,
        }
    )
    placeholder_map = {
        placeholder: value
        for placeholder, value in placeholder_map.items()
        if not _is_block_placeholder(placeholder)
    }

    measure_catalog = None
    if any(
        ph in ("{MEASURE_BLOCK}", "{MEASURE_SUMMARY_ROW}") for ph in block_placeholders
    ):
        try:
            measure_catalog = load_measure_catalog()
        except FileNotFoundError:
            measure_catalog = None

    blocks_rendered: List[str] = []
    blocks_unresolved: List[str] = []
    block_replacements: Dict[str, str] = {}
    for placeholder in block_placeholders:
        renderer = get_block_renderer(placeholder)
        if renderer:
            renderer_kwargs: Dict[str, Any] = {"schema": None, "mapping": mapping_data}
            if placeholder in ("{MEASURE_BLOCK}", "{MEASURE_SUMMARY_ROW}"):
                renderer_kwargs.update(
                    {
                        "catalog": measure_catalog,
                        "placeholders": base_placeholder_map,
                    }
                )
            rendered_text = renderer(project_data, **renderer_kwargs)
        else:
            rendered_text = DEFAULT_EMPTY_BLOCK_TEXT
        if not _has_meaningful_value(rendered_text):
            rendered_text = DEFAULT_EMPTY_BLOCK_TEXT
            blocks_unresolved.append(placeholder)
        block_replacements[placeholder] = rendered_text
        blocks_rendered.append(placeholder)

    replacement_map = {**placeholder_map, **block_replacements}
    block_paragraphs = {
        placeholder: _split_block_paragraphs(text)
        for placeholder, text in block_replacements.items()
    }

    placeholders_found = len(placeholder_occurrences)
    placeholders_replaced = 0
    unresolved: Set[str] = set()

    measure_styles = _ensure_measure_styles(doc)
    structured_measures = measure_narratives.collect_structured_measures(project_data)
    measure_fallback_text = block_replacements.get("{MEASURE_BLOCK}", DEFAULT_EMPTY_BLOCK_TEXT)
    if "{MEASURE_SUMMARY_ROW}" in block_placeholders:
        _fill_measure_summary_table(doc, project_data, structured_measures, measure_catalog)

    for paragraph in _iter_all_paragraphs(doc):
        text = paragraph.text
        if not text or "{" not in text:
            continue
        placeholders_replaced += _replace_placeholders_in_runs(paragraph, placeholder_map)
        text = paragraph.text
        found = PLACEHOLDER_PATTERN.findall(text)
        if not found:
            continue
        replaced_text = text
        expanded_block = False
        for placeholder in set(found):
            if placeholder in replacement_map:
                if (
                    placeholder == "{MEASURE_BLOCK}"
                    and text.strip() == placeholder
                ):
                    _insert_measure_block(
                        paragraph,
                        structured_measures,
                        measure_styles,
                        measure_fallback_text,
                    )
                    expanded_block = True
                    continue
                if placeholder in block_replacements:
                    placeholders_replaced += text.count(placeholder)
                    paragraphs = block_paragraphs.get(placeholder, [])
                    if len(paragraphs) > 1:
                        if text.strip() == placeholder:
                            paragraph.text = paragraphs[0]
                            current_para = paragraph
                            for block_text in paragraphs[1:]:
                                current_para = _add_paragraph_after(
                                    current_para, block_text, style=paragraph.style
                                )
                            expanded_block = True
                            continue
                        prefix, suffix = text.split(placeholder, 1)
                        paragraph.text = f"{prefix}{paragraphs[0]}{suffix}"
                        current_para = paragraph
                        for block_text in paragraphs[1:]:
                            current_para = _add_paragraph_after(
                                current_para, block_text, style=paragraph.style
                            )
                        expanded_block = True
                        continue
                    joined = " ".join(paragraphs) if paragraphs else replacement_map[placeholder]
                    replaced_text = replaced_text.replace(placeholder, joined)
                else:
                    placeholders_replaced += _replace_placeholders_in_runs(
                        paragraph, {placeholder: replacement_map[placeholder]}
                    )
                    text = paragraph.text
            else:
                unresolved.add(placeholder)
        if expanded_block:
            continue
        if replaced_text != text:
            paragraph.text = replaced_text

    output_dir = os.path.dirname(out_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
    doc.save(out_path)

    placeholders_replaced += _replace_placeholders_in_docx_xml(out_path, replacement_map)

    remaining_placeholders = _collect_placeholders_from_docx(out_path)
    unresolved.update(remaining_placeholders)

    summary = {
        "out_path": out_path,
        "placeholders_found": placeholders_found,
        "placeholders_replaced": placeholders_replaced,
        "unresolved": sorted(unresolved),
        "blocks_rendered": blocks_rendered,
        "blocks_unresolved": blocks_unresolved,
    }
    if any(
        ph in ("{MEASURE_BLOCK}", "{MEASURE_SUMMARY_ROW}") for ph in block_placeholders
    ):
        summary["measures_rendered_count"] = measure_narratives.count_selected_measures(
            project_data, catalog=measure_catalog
        )

    if strict and unresolved:
        summary["strict_error"] = True

    return summary
